from docx import Document
from typing import Dict, Any, Optional
import re


class DOCXParser:
    """Parser for DOCX resume files."""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"DOCX parsing error: {str(e)}")
    
    @staticmethod
    def parse_resume(text: str) -> Dict[str, Any]:
        """Parse resume text into structured data."""
        # Reuse the same parsing logic as PDF
        from app.parsers.pdf_parser import PDFParser
        return PDFParser.parse_resume(text)
